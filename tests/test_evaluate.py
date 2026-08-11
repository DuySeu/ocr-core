import pytest
import yaml
from docx import Document

from evaluate import (
    NO_GOLD_REASON,
    NO_PREDICTION_REASON,
    EvalConfig,
    PairingError,
    evaluate_engine,
    write_report,
)
from evaluate.config import ConfigError, load_config
from evaluate.engines import UnknownEngineError
from evaluate.report import render_markdown

TEXT_BOX = {"id": 10001, "category": "text", "bbox": (10, 20, 40, 30), "text": "Điều 1"}


# Build an output directory and a ground-truth directory, and the config joining them.
@pytest.fixture
def workspace(tmp_path):
    def _build(engine="tesseract", predictions=None, ground_truths=None):
        output_dir = tmp_path / "output"
        ground_truth_dir = tmp_path / "ground_truth"
        output_dir.mkdir(exist_ok=True)
        ground_truth_dir.mkdir(exist_ok=True)

        for name, body in (predictions or {}).items():
            (output_dir / name).write_text(body, encoding="utf-8")
        for name, body in (ground_truths or {}).items():
            (ground_truth_dir / name).write_text(body, encoding="utf-8")

        return EvalConfig(
            engine=engine, output_dir=output_dir, ground_truth_dir=ground_truth_dir
        )

    return _build


def test_a_prediction_is_paired_with_the_ground_truth_of_the_same_stem(workspace):
    config = workspace(
        predictions={"doc.md": "Điều 1. Phạm vi"},
        ground_truths={"doc.md": "Điều 1. Phạm vi"},
    )

    report = evaluate_engine(config)

    assert [d.doc_id for d in report.documents] == ["doc"]
    assert report.documents[0].text.cer == 0.0
    assert report.corpus_text.cer == 0.0


def test_a_prediction_with_no_ground_truth_is_named_rather_than_dropped(workspace):
    config = workspace(predictions={"doc.md": "Điều 1"}, ground_truths={"other.md": "x"})

    report = evaluate_engine(config)

    assert report.documents[0].text is None
    assert "no ground-truth .md/.docx" in report.documents[0].notes[0]


def test_ground_truth_with_no_prediction_is_listed_not_ignored(workspace):
    config = workspace(predictions={"doc.md": "Điều 1"}, ground_truths={"lonely.md": "x"})

    report = evaluate_engine(config)

    assert ("lonely", NO_PREDICTION_REASON) in report.unpaired


def test_a_prediction_with_no_ground_truth_is_named_in_the_unpaired_list_too(workspace):
    config = workspace(predictions={"doc.md": "Điều 1"}, ground_truths={"lonely.md": "x"})

    report = evaluate_engine(config)

    assert ("doc", NO_GOLD_REASON) in report.unpaired


def test_two_predictions_sharing_a_stem_is_an_error_not_a_silent_pick(workspace, tmp_path):
    config = workspace(predictions={"doc.md": "Điều 1"}, ground_truths={"doc.md": "Điều 1"})
    nested = tmp_path / "output" / "nested"
    nested.mkdir()
    (nested / "doc.md").write_text("Điều 1", encoding="utf-8")

    with pytest.raises(PairingError, match="share the stem"):
        evaluate_engine(config)


def test_the_corpus_row_pools_characters_rather_than_averaging_document_rates(workspace):
    # One character wrong in a short document and none in a long one
    config = workspace(
        predictions={"short.md": "abx", "long.md": "y" * 100},
        ground_truths={"short.md": "abc", "long.md": "y" * 100},
    )

    report = evaluate_engine(config)

    assert report.corpus_text.cer == pytest.approx(1 / 103)


def test_chandra_output_reports_why_its_json_cannot_give_boxes(tmp_path):
    output_dir = tmp_path / "output"
    ground_truth_dir = tmp_path / "gt"
    output_dir.mkdir()
    ground_truth_dir.mkdir()
    (output_dir / "doc.md").write_text("Điều 1", encoding="utf-8")
    (output_dir / "doc_metadata.json").write_text('{"pages": []}', encoding="utf-8")
    (ground_truth_dir / "doc.md").write_text("Điều 1", encoding="utf-8")

    report = evaluate_engine(
        EvalConfig(engine="chandra", output_dir=output_dir, ground_truth_dir=ground_truth_dir)
    )

    assert report.documents[0].layout is None
    assert "no per-element bbox" in report.documents[0].notes[0]


def test_boxes_on_both_sides_are_matched_and_scored(workspace, write_coco, tmp_path):
    config = workspace(predictions={"doc.md": "Điều 1"}, ground_truths={"doc.md": "Điều 1"})
    write_coco("output/doc.coco.json", [TEXT_BOX])
    write_coco("ground_truth/doc.coco.json", [TEXT_BOX])

    report = evaluate_engine(config)

    layout = report.documents[0].layout
    assert layout.pages_scored == 1
    assert layout.categories[0].precision == 1.0
    assert layout.element_text.cer == 0.0


def test_a_page_the_pipeline_failed_on_is_not_counted_as_a_detector_miss(
    workspace, write_coco
):
    config = workspace(predictions={"doc.md": "x"}, ground_truths={"doc.md": "x"})
    write_coco(
        "ground_truth/doc.coco.json",
        [TEXT_BOX, {**TEXT_BOX, "id": 20001, "page": 2}],
        pages=(1, 2),
    )
    write_coco("output/doc.coco.json", [TEXT_BOX], pages=(1,), page_errors=[2])

    report = evaluate_engine(config)

    layout = report.documents[0].layout
    assert layout.pages_scored == 1
    assert "page 2" in layout.page_notes[0]


def test_a_deskew_mismatch_makes_the_page_unscoreable_rather_than_scoring_it_wrong(
    workspace, write_coco
):
    config = workspace(predictions={"doc.md": "x"}, ground_truths={"doc.md": "x"})
    write_coco("ground_truth/doc.coco.json", [TEXT_BOX], deskew_angle=1.4)
    write_coco("output/doc.coco.json", [TEXT_BOX], deskew_angle=0.0)

    report = evaluate_engine(config)

    assert report.documents[0].layout.pages_scored == 0
    assert "coordinate frames differ" in report.documents[0].layout.page_notes[0]


def test_predictions_on_pages_absent_from_gold_are_skipped_not_counted_as_spurious(
    workspace, write_coco
):
    config = workspace(predictions={"doc.md": "x"}, ground_truths={"doc.md": "x"})
    write_coco("ground_truth/doc.coco.json", [TEXT_BOX], pages=(1,))
    write_coco(
        "output/doc.coco.json",
        [TEXT_BOX, {**TEXT_BOX, "id": 20001, "page": 2}],
        pages=(1, 2),
    )

    report = evaluate_engine(config)

    layout = report.documents[0].layout
    assert layout.pages_scored == 1
    assert layout.categories[0].n_pred == 1


def test_an_engine_with_no_adapter_fails_by_name_rather_than_guessing(workspace):
    config = workspace(engine="nonesuch", predictions={"doc.md": "x"})

    with pytest.raises(UnknownEngineError, match="nonesuch"):
        evaluate_engine(config)


def test_the_report_file_is_named_after_the_output_dir(workspace, tmp_path):
    config = workspace(predictions={"doc.md": "Điều 1"}, ground_truths={"doc.md": "Điều 1"})
    results_dir = tmp_path / "results"

    report_path = write_report(evaluate_engine(config), results_dir)

    assert report_path.name == "output_results.md"
    assert report_path.parent == results_dir


def test_config_reads_only_the_three_keys_scoring_needs(tmp_path):
    (tmp_path / "output").mkdir()
    (tmp_path / "gt").mkdir()
    path = tmp_path / "config.yaml"
    path.write_text(
        yaml.safe_dump(
            {
                "engine": "chandra",
                "langs": ["vie"],
                "preprocess_steps": ["deskew"],
                "output_dir": str(tmp_path / "output"),
                "ground_truth_dir": str(tmp_path / "gt"),
            }
        ),
        encoding="utf-8",
    )

    config = load_config(path)

    assert config.engine == "chandra"
    assert config.results_dir.name == "chandra"


def test_a_config_missing_ground_truth_dir_says_which_key(tmp_path):
    path = tmp_path / "config.yaml"
    path.write_text(yaml.safe_dump({"engine": "chandra", "output_dir": "."}), encoding="utf-8")

    with pytest.raises(ConfigError, match="ground_truth_dir"):
        load_config(path)


def test_every_section_renders_even_when_nothing_could_be_scored(workspace):
    config = workspace(predictions={"doc.md": "Điều 1"}, ground_truths={})

    markdown = render_markdown(evaluate_engine(config))

    for heading in (
        "## 1 · Text",
        "## 2 · Layout",
        "## 3 · Text",
        "## 4 · Tables",
        "## 5 · Not measured",
    ):
        assert heading in markdown


# ---------- tables ----------

TABLE_HTML = "<table><tr><td>An</td><td>10</td></tr></table>"
PIPE_TABLE = "| Tên | Số |\n| --- | --- |\n| An | 10 |"


def test_tables_are_paired_and_scored_at_document_level(workspace):
    config = workspace(
        predictions={"doc.md": f"Mở đầu\n\n{TABLE_HTML}"},
        ground_truths={"doc.md": f"Mở đầu\n\n{TABLE_HTML}"},
    )

    tables = evaluate_engine(config).documents[0].tables

    assert tables.n_matched == 1
    assert tables.teds == pytest.approx(1.0)
    assert tables.table_recall == pytest.approx(1.0)


def test_a_pipe_gold_table_is_paired_with_an_html_predicted_table(workspace):
    config = workspace(
        predictions={"doc.md": "<table><tr><th>Tên</th><th>Số</th></tr>"
                               "<tr><td>An</td><td>10</td></tr></table>"},
        ground_truths={"doc.md": PIPE_TABLE},
    )

    tables = evaluate_engine(config).documents[0].tables

    assert tables.n_matched == 1
    assert tables.teds == pytest.approx(1.0)


def test_a_document_with_no_ground_truth_file_is_not_scoreable_rather_than_empty(workspace):
    config = workspace(predictions={"doc.md": TABLE_HTML}, ground_truths={})

    document = evaluate_engine(config).documents[0]

    assert document.tables is None
    assert "not scoreable - no ground truth" in render_markdown(evaluate_engine(config))


def test_gold_with_no_tables_reports_the_predicted_count_not_a_recall(workspace):
    config = workspace(
        predictions={"doc.md": TABLE_HTML}, ground_truths={"doc.md": "Không có bảng"}
    )

    tables = evaluate_engine(config).documents[0].tables

    assert tables.n_gold == 0
    assert tables.n_pred == 1
    assert tables.table_recall is None


def test_gold_tables_are_read_from_a_docx_by_suffix(workspace, tmp_path):
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "An"
    table.rows[0].cells[1].text = "10"
    config = workspace(predictions={"doc.md": TABLE_HTML})
    document.save(str(tmp_path / "ground_truth" / "doc.docx"))

    tables = evaluate_engine(config).documents[0].tables

    assert tables.n_gold == 1
    assert tables.n_matched == 1


def test_corpus_tables_are_weighted_by_matched_count_not_averaged_per_document(workspace):
    wrong_cell = "<table><tr><td>An</td><td>11</td></tr></table>"
    two_tables = f"{TABLE_HTML}\n\n{TABLE_HTML}"
    config = workspace(
        predictions={"one.md": wrong_cell, "two.md": two_tables},
        ground_truths={"one.md": TABLE_HTML, "two.md": two_tables},
    )

    report = evaluate_engine(config)

    # One imperfect table against two perfect ones, pooled over three matches
    per_document = [d.tables.teds for d in report.documents]
    assert report.corpus_tables.n_matched == 3
    assert report.corpus_tables.teds != pytest.approx(sum(per_document) / 2)
    assert report.corpus_tables.teds == pytest.approx(
        sum(d.tables.teds * d.tables.n_matched for d in report.documents) / 3
    )


def test_the_pairing_floor_is_printed_in_the_table_section_heading(workspace):
    config = workspace(predictions={"doc.md": "x"}, ground_truths={"doc.md": "x"})

    markdown = render_markdown(evaluate_engine(config))

    assert "## 4 · Tables  (pairing floor >= 0.5" in markdown


def test_a_document_with_no_table_on_either_side_renders_a_hyphen_not_n_a(workspace):
    config = workspace(predictions={"doc.md": "Không bảng"}, ground_truths={"doc.md": "Không bảng"})

    markdown = render_markdown(evaluate_engine(config))

    assert "| doc | - | - | 0 | 0 | 0 | - | |" in markdown
