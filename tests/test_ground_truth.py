import pytest
from docx import Document

from evaluate.ground_truth import GroundTruthError, discover_boxes, discover_text, load


def test_markdown_ground_truth_is_read_verbatim(tmp_path):
    path = tmp_path / "gt.md"
    path.write_text("## Tiêu đề\n", encoding="utf-8")

    assert load(path) == "## Tiêu đề\n"


def test_docx_table_cells_are_read_in_document_order_with_the_paragraphs(tmp_path):
    document = Document()
    document.add_paragraph("Điều 1. Phạm vi")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "An"
    table.rows[0].cells[1].text = "10"
    document.add_paragraph("Điều 2. Giải thích")
    path = tmp_path / "gt.docx"
    document.save(str(path))

    assert load(path).split("\n") == ["Điều 1. Phạm vi", "An", "10", "Điều 2. Giải thích"]


def test_an_unreadable_format_is_named_rather_than_read_as_bytes(tmp_path):
    path = tmp_path / "gt.pdf"
    path.write_bytes(b"%PDF-1.4")

    with pytest.raises(GroundTruthError, match="readable formats"):
        load(path)


def test_discovery_keys_files_by_stem_across_both_formats(tmp_path):
    (tmp_path / "a.md").write_text("a", encoding="utf-8")
    (tmp_path / "nested").mkdir()
    Document().save(str(tmp_path / "nested" / "b.docx"))

    assert sorted(discover_text(tmp_path)) == ["a", "b"]


def test_two_ground_truth_files_sharing_a_stem_is_an_error_not_a_silent_pick(tmp_path):
    (tmp_path / "a.md").write_text("a", encoding="utf-8")
    Document().save(str(tmp_path / "a.docx"))

    with pytest.raises(GroundTruthError, match="share the stem"):
        discover_text(tmp_path)


def test_a_pdf_beside_the_ground_truth_is_ignored_not_treated_as_a_target(tmp_path):
    (tmp_path / "a.md").write_text("a", encoding="utf-8")
    (tmp_path / "a-scan.pdf").write_bytes(b"%PDF-1.4")

    assert sorted(discover_text(tmp_path)) == ["a"]


def test_box_files_key_on_the_stem_with_the_compound_suffix_stripped(tmp_path):
    (tmp_path / "a.coco.json").write_text("{}", encoding="utf-8")
    (tmp_path / "b.json").write_text("{}", encoding="utf-8")

    assert sorted(discover_boxes(tmp_path)) == ["a", "b"]
