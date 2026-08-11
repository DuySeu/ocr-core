import pytest
from docx import Document

from evaluate.table_extract import (
    TableExtractError,
    extract_docx_tables,
    extract_html_tables,
    walk_docx_cells,
)

EMBEDDED = "<table><tr><td>a</td><td>b</td></tr></table>"
NESTED_OUTER = "<table><tr><td><table><tr><td>inner</td></tr></table></td></tr></table>"

PIPE_TABLE = """| TT | Ngày |
| --- | ----- |
| 1 | 19/03 |
| 2 | 21/03 |"""


# Build a docx in tmp_path with one table, applying merges by (r0, c0, r1, c1) spans.
@pytest.fixture
def docx_with_table(tmp_path):
    def _build(rows, cols, merges=(), name="t.docx"):
        document = Document()
        table = document.add_table(rows=rows, cols=cols)
        for row in range(rows):
            for column in range(cols):
                table.cell(row, column).text = f"r{row}c{column}"
        for r0, c0, r1, c1 in merges:
            table.cell(r0, c0).merge(table.cell(r1, c1))

        path = tmp_path / name
        document.save(str(path))
        return path

    return _build


# ---------- extract_html_tables ----------


def test_extracts_every_table_from_markdown_with_embedded_html():
    text = f"Mở đầu\n\n{EMBEDDED}\n\nGiữa\n\n{EMBEDDED}\n\nKết"

    assert len(extract_html_tables(text)) == 2


def test_ignores_a_table_nested_inside_another_table():
    found = extract_html_tables(NESTED_OUTER)

    assert len(found) == 1
    assert "inner" in found[0]


def test_returns_an_empty_list_for_a_document_with_no_tables():
    assert extract_html_tables("Chỉ có văn bản thường, không bảng nào.") == []


@pytest.mark.parametrize("text", ["", "   ", "\n\n", "\t"])
def test_returns_an_empty_list_for_empty_and_whitespace_only_input(text):
    assert extract_html_tables(text) == []


def test_converts_a_markdown_pipe_table_to_html():
    found = extract_html_tables(PIPE_TABLE)

    assert len(found) == 1
    assert found[0] == (
        "<table><tr><th>TT</th><th>Ngày</th></tr>"
        "<tr><td>1</td><td>19/03</td></tr>"
        "<tr><td>2</td><td>21/03</td></tr></table>"
    )


def test_does_not_treat_a_pipe_run_without_a_delimiter_row_as_a_table():
    assert extract_html_tables("| a | b |\n| c | d |") == []


def test_does_not_start_a_table_from_a_bullet_line_containing_a_pipe():
    # This exact shape appears in four lpbank ground-truth files
    text = "- BKS; | (Để báo cáo)\n- Chủ tịch HĐQT; | Để báo cáo"

    assert extract_html_tables(text) == []


def test_keeps_an_escaped_pipe_as_literal_text():
    found = extract_html_tables("| a\\|b | c |\n| --- | --- |\n| d | e |")

    assert "<th>a|b</th>" in found[0]


def test_reads_both_pipe_tables_and_embedded_html_from_one_document():
    assert len(extract_html_tables(f"{PIPE_TABLE}\n\n{EMBEDDED}")) == 2


def test_returns_tables_in_document_order_when_both_forms_are_present():
    pipe_first = extract_html_tables(f"{PIPE_TABLE}\n\n{EMBEDDED}")
    html_first = extract_html_tables(f"{EMBEDDED}\n\n{PIPE_TABLE}")

    assert "<th>TT</th>" in pipe_first[0] and "<td>a</td>" in pipe_first[1]
    assert "<td>a</td>" in html_first[0] and "<th>TT</th>" in html_first[1]


def test_escapes_markup_characters_coming_from_a_pipe_cell():
    found = extract_html_tables("| a < b | c |\n| --- | --- |\n| d | e |")

    assert "&lt;" in found[0]


# ---------- extract_docx_tables ----------


def test_converts_a_docx_table_to_html(docx_with_table):
    found = extract_docx_tables(docx_with_table(rows=2, cols=2))

    assert len(found) == 1
    assert found[0] == (
        "<table><tr><td>r0c0</td><td>r0c1</td></tr>"
        "<tr><td>r1c0</td><td>r1c1</td></tr></table>"
    )


def test_emits_colspan_from_grid_span(docx_with_table):
    # Merging two cells in one row collapses them into a single gridSpan cell
    found = extract_docx_tables(docx_with_table(rows=2, cols=3, merges=[(0, 0, 0, 1)]))

    assert 'colspan="2"' in found[0]


def test_emits_rowspan_from_v_merge_and_no_cell_for_continuation_rows(docx_with_table):
    found = extract_docx_tables(docx_with_table(rows=2, cols=2, merges=[(0, 0, 1, 0)]))

    assert 'rowspan="2"' in found[0]
    # Row 1 keeps only its unmerged cell
    assert found[0].endswith("<tr><td>r1c1</td></tr></table>")


def test_emits_a_cross_row_merged_cell_once(docx_with_table):
    found = extract_docx_tables(docx_with_table(rows=2, cols=2, merges=[(0, 0, 1, 0)]))

    assert found[0].count("r0c0") == 1


def test_raises_when_the_docx_does_not_exist(tmp_path):
    with pytest.raises(TableExtractError, match="docx not found"):
        extract_docx_tables(tmp_path / "absent.docx")


# ---------- walk_docx_cells ----------


def test_walks_cells_row_major(docx_with_table):
    document = Document(str(docx_with_table(rows=2, cols=2)))

    positions = [(c.row, c.column) for c in walk_docx_cells(document.tables[0])]

    assert positions == [(0, 0), (0, 1), (1, 0), (1, 1)]


def test_treats_a_cell_with_no_v_merge_as_unmerged_not_as_a_continuation(docx_with_table):
    document = Document(str(docx_with_table(rows=2, cols=2)))

    cells = list(walk_docx_cells(document.tables[0]))

    assert len(cells) == 4
    assert {c.rowspan for c in cells} == {1}


def test_advances_the_grid_column_by_grid_span_so_later_columns_are_not_shifted(
    docx_with_table,
):
    document = Document(str(docx_with_table(rows=2, cols=3, merges=[(0, 0, 0, 1)])))

    row_zero = [c for c in walk_docx_cells(document.tables[0]) if c.row == 0]

    assert [(c.column, c.colspan) for c in row_zero] == [(0, 2), (2, 1)]


def test_joins_paragraphs_within_a_cell_rather_than_gluing_them(docx_with_table):
    path = docx_with_table(rows=2, cols=2, merges=[(0, 0, 1, 0)])
    document = Document(str(path))

    merged = next(c for c in walk_docx_cells(document.tables[0]) if c.rowspan == 2)

    assert merged.text == "r0c0\nr1c0"
